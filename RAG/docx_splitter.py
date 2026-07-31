"""DOCX document splitter built on python-docx.

Produces exactly the same chunk contract as `md_splitter` (see `chunking`), so
the ClickHouse schema and all knowledge-base tools work unchanged for a .docx
corpus.

Structure recovery:
  - Section breadcrumb comes from heading styles: "Heading N" / "Заголовок N"
    (Word stores localised style names) and, as a fallback, from the style's
    `w:outlineLvl` — that catches custom heading styles.
  - Consecutive body paragraphs are merged into one prose block until a heading,
    a table, or `chunk_size` is reached. Word paragraphs are often a single
    sentence or list item, so per-paragraph chunks would be too fine-grained.
  - Tables are read from the Word object model: first row = headers, the rest =
    data rows. Each table yields a `table_full` doc plus one `table_row` per row.

Positions: .docx has no line numbers, so `line_start`/`line_end` hold 1-based
ordinals of the body blocks the chunk spans. That keeps ordering and
`get_neighbor_chunks` meaningful.
"""
from __future__ import annotations

import io
import json
import logging
import posixpath
import re
import zipfile
from collections.abc import Iterator
from pathlib import Path
from typing import Optional, Union
from xml.etree import ElementTree

from docx import Document as open_docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph
from langchain_core.documents import Document

from chunking import (
    ChunkIndexer,
    SectionStack,
    clean_text,
    normalize_headers,
    prose_to_documents,
    table_rows_to_documents,
)

logger = logging.getLogger(__name__)

# Heading styles. Beyond stock "Heading N" / "Заголовок N", corporate templates
# prefix them ("Приложение: Заголовок 2"), so the marker is matched anywhere in the
# name. The lookbehind keeps "Подзаголовок" (a caption style) from matching.
_HEADING_LEVEL_RE = re.compile(r"(?<![a-zа-яё])(?:heading|заголовок)\s*(\d+)", re.IGNORECASE)
_HEADING_WORD_RE  = re.compile(r"(?<![a-zа-яё])(?:heading|заголовок)(?![a-zа-яё])", re.IGNORECASE)
_TITLE_STYLE_RE   = re.compile(r"^(?:title|название)$", re.IGNORECASE)
_LIST_STYLE_RE    = re.compile(r"(list|список|bullet|маркирован|нумерован)", re.IGNORECASE)

# Navigation headings: they open a level in the document but carry no content of
# their own, so keeping them would prefix every following breadcrumb.
_SERVICE_HEADING_RE = re.compile(
    r"^(?:содержание|оглавление|перечень\s+таблиц|перечень\s+рисунков"
    r"|table\s+of\s+contents|contents)$",
    re.IGNORECASE,
)

_MAX_HEADING_LEVEL = 9

# Максимальная длина ячейки, которую ещё можно принять за название колонки.
_MAX_HEADER_CELL_CHARS = 80

BlockItem = Union[Paragraph, Table]


# Relationship namespace of the OPC package (.docx is a zip of OPC parts).
_RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


# ---------------------------------------------------------------------------
# Opening the package (with repair for dangling relationships)
# ---------------------------------------------------------------------------

def _prune_dangling_rels(rels_xml: bytes, part_dir: str, entries: set[str]) -> tuple[bytes, int]:
    """Drop internal relationships whose target is missing from the archive.

    Returns the rewritten XML and the number of relationships removed.
    """
    ElementTree.register_namespace("", _RELS_NS)
    root = ElementTree.fromstring(rels_xml)

    removed = 0
    for relationship in list(root):
        if relationship.get("TargetMode") == "External":
            continue
        target = relationship.get("Target") or ""
        resolved = posixpath.normpath(posixpath.join(part_dir, target)).lstrip("/")
        if resolved not in entries:
            root.remove(relationship)
            removed += 1

    return ElementTree.tostring(root, encoding="UTF-8", xml_declaration=True), removed


def _repair_package(docx_file: Path) -> io.BytesIO:
    """Rebuild the .docx package in memory without dangling relationships.

    Word tolerates relationships pointing at missing parts (`Target="NULL"` for a
    dropped image, for example); python-docx loads every part eagerly and fails.
    The source file is never modified.
    """
    buffer = io.BytesIO()
    total_removed = 0

    with zipfile.ZipFile(docx_file) as source:
        entries = set(source.namelist())
        with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as target:
            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename.endswith(".rels"):
                    part_dir = posixpath.dirname(posixpath.dirname(item.filename))
                    try:
                        data, removed = _prune_dangling_rels(data, part_dir, entries)
                        total_removed += removed
                    except ElementTree.ParseError as exc:
                        logger.warning(f"[{docx_file.name}] не разобран {item.filename}: {exc}")
                target.writestr(item, data)

    logger.info(f"[{docx_file.name}] пакет восстановлен, удалено битых связей: {total_removed}")
    buffer.seek(0)
    return buffer


def open_document(docx_file: Path) -> DocxDocument:
    """Open a .docx, repairing dangling relationships if the direct load fails."""
    try:
        return open_docx(str(docx_file))
    except KeyError as exc:
        logger.warning(f"[{docx_file.name}] битая структура пакета ({exc}), пробуем восстановить")
        return open_docx(_repair_package(docx_file))


# ---------------------------------------------------------------------------
# Body traversal
# ---------------------------------------------------------------------------

def _iter_block_items(parent: DocxDocument | _Cell) -> Iterator[BlockItem]:
    """Yield paragraphs and tables of `parent` in document order.

    python-docx exposes `.paragraphs` and `.tables` as separate collections,
    which loses their relative order — the XML children must be walked directly.
    """
    if isinstance(parent, _Cell):
        parent_element = parent._tc
    else:
        parent_element = parent.element.body

    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# ---------------------------------------------------------------------------
# Paragraph classification
# ---------------------------------------------------------------------------

def _style_name(paragraph: Paragraph) -> str:
    """Style name of the paragraph, or "" when the style is missing."""
    try:
        style = paragraph.style
        return style.name or "" if style is not None else ""
    except (AttributeError, KeyError):
        return ""


def _outline_level(paragraph: Paragraph) -> int | None:
    """Heading level from `w:outlineLvl` (0-based in XML), if declared.

    Checks the paragraph's direct formatting first, then its style definition —
    this is what makes custom heading styles ("Пункт 1.1") resolve correctly.
    """
    for properties in (paragraph._p.pPr, getattr(paragraph.style, "element", None)):
        if properties is None:
            continue
        holder = properties if properties.tag == qn("w:pPr") else properties.find(qn("w:pPr"))
        if holder is None:
            continue
        outline = holder.find(qn("w:outlineLvl"))
        if outline is not None:
            raw = outline.get(qn("w:val"))
            if raw is not None and raw.isdigit():
                level = int(raw) + 1
                if 1 <= level <= _MAX_HEADING_LEVEL:
                    return level
    return None


def _heading_level(paragraph: Paragraph) -> int | None:
    """Return the heading level of a paragraph, or None if it is body text.

    Resolution order: numbered heading style ("Заголовок 2", "Приложение: Заголовок 2")
    -> document title -> unnumbered heading style ("Заголовок: Технический") ->
    the style's declared outline level.
    """
    name = _style_name(paragraph)

    match = _HEADING_LEVEL_RE.search(name)
    if match:
        level = int(match.group(1))
        if 1 <= level <= _MAX_HEADING_LEVEL:
            return level

    if _TITLE_STYLE_RE.match(name) or _HEADING_WORD_RE.search(name):
        return 1

    return _outline_level(paragraph)


def _is_list_item(paragraph: Paragraph) -> bool:
    """True if the paragraph is a bullet/numbered list item."""
    if _LIST_STYLE_RE.search(_style_name(paragraph)):
        return True
    properties = paragraph._p.pPr
    return properties is not None and properties.find(qn("w:numPr")) is not None


def _paragraph_text(paragraph: Paragraph) -> str:
    """Visible text of a paragraph, list items prefixed with a dash."""
    text = clean_text(paragraph.text)
    if not text:
        return ""
    return f"- {text}" if _is_list_item(paragraph) else text


# ---------------------------------------------------------------------------
# Table parsing
# ---------------------------------------------------------------------------

def _cell_text(cell: _Cell) -> str:
    """Flatten a table cell (including nested tables) into a single line."""
    parts: list[str] = []
    for block in _iter_block_items(cell):
        if isinstance(block, Paragraph):
            text = clean_text(block.text)
            if text:
                parts.append(text)
        else:
            for nested_row in block.rows:
                nested = " | ".join(_cell_text(c) for c in nested_row.cells)
                if nested.strip(" |"):
                    parts.append(nested)
    return " ".join(parts)


def _distinct_cells(row: list[str]) -> int:
    """Number of distinct non-empty values in a row."""
    return len({cell.strip() for cell in row if cell.strip()})


def _looks_like_header(row: list[str]) -> bool:
    """Whether a row can plausibly serve as the header row.

    Column names are short and distinct. A row with a long cell is data — such a
    cell must never become a column name.
    """
    if _distinct_cells(row) < 2:
        return False
    return all(len(cell.strip()) <= _MAX_HEADER_CELL_CHARS for cell in row)


def _marked_as_header(row: _Row) -> bool:
    """True if Word marks the row as a repeating header (`w:tblHeader`)."""
    properties = row._tr.trPr
    return properties is not None and properties.find(qn("w:tblHeader")) is not None


def _cell_is_bold(cell: _Cell) -> bool:
    """True if all visible text of the cell is bold (directly or via its style)."""
    paragraphs = [p for p in cell.paragraphs if p.text.strip()]
    if not paragraphs:
        return False

    runs = [run for p in paragraphs for run in p.runs if run.text.strip()]
    if runs and all(run.bold for run in runs):
        return True
    return all(p.style is not None and p.style.font.bold for p in paragraphs)


def _cell_is_shaded(cell: _Cell) -> bool:
    """True if the cell has a background fill (a common header decoration)."""
    properties = cell._tc.tcPr
    if properties is None:
        return False
    shading = properties.find(qn("w:shd"))
    if shading is None:
        return False
    fill = shading.get(qn("w:fill")) or ""
    return bool(fill) and fill.lower() not in ("auto", "ffffff")


def _has_header_formatting(row: _Row) -> bool:
    """Whether the row is visually marked up as a header row.

    Word almost always distinguishes a header row: the "repeat as header" flag,
    bold text or a background fill. Position alone is not evidence — glossary and
    reference tables often start with data right away.
    """
    if _marked_as_header(row):
        return True
    cells = [cell for cell in row.cells if cell.text.strip()]
    if not cells:
        return False
    return (all(_cell_is_bold(cell) for cell in cells)
            or all(_cell_is_shaded(cell) for cell in cells))


def _find_header_row(table: Table, rows: list[list[str]]) -> Optional[int]:
    """Index of the header row, or None if the table has no header row.

    A row qualifies only on visual evidence (see `_has_header_formatting`). When
    the first row is a single merged caption spanning the width, the second row is
    considered instead. Tables without evidence — glossaries and reference lists
    that start with data — get synthetic column names, so no data row is consumed.
    """
    table_rows = table.rows

    if (len(rows) >= 3 and _distinct_cells(rows[0]) <= 1
            and (_has_header_formatting(table_rows[1]) or _looks_like_header(rows[1]))):
        return 1

    if _has_header_formatting(table_rows[0]) and any(cell.strip() for cell in rows[0]):
        return 0

    return None


def _parse_table(table: Table) -> tuple[list[str], list[list[str]]]:
    """Extract (headers, data_rows) from a Word table.

    Merged cells repeat their text in `row.cells`, so all rows keep equal width.
    Headers of a header-less table are synthesised as empty strings and get their
    `Колонка N` names in `chunking.normalize_headers`.
    """
    rows = [[_cell_text(cell) for cell in row.cells] for row in table.rows]
    if not rows:
        return [], []

    header_index = _find_header_row(table, rows)

    if header_index is None:
        width = max(len(row) for row in rows)
        headers = [""] * width
        data_rows = rows
    else:
        headers = rows[header_index]
        data_rows = rows[header_index + 1:]

    data_rows = [row for row in data_rows if any(cell.strip() for cell in row)]
    if not data_rows:
        return [], []
    return headers, data_rows


def _render_table_text(headers: list[str], data_rows: list[list[str]]) -> str:
    """Render a table as a pipe-style block for the `table_full` chunk."""
    lines = ["| " + " | ".join(headers) + " |",
             "| " + " | ".join("---" for _ in headers) + " |"]
    for row in data_rows:
        padded = (row + [""] * max(0, len(headers) - len(row)))[: len(headers)]
        lines.append("| " + " | ".join(padded) + " |")
    return "\n".join(lines)


def _raw_table_text(table: Table) -> str:
    """Fallback rendering for tables without a usable header row."""
    return "\n".join(
        " | ".join(_cell_text(cell) for cell in row.cells) for row in table.rows
    ).strip()


# ---------------------------------------------------------------------------
# Main public API
# ---------------------------------------------------------------------------

def split_docx_file(
    docx_file: Path,
    chunk_size: int = 1500,
    chunk_overlap: int = 150,
) -> list[Document]:
    """Parse a .docx file into LangChain Documents.

    Args:
        docx_file:     Path to the .docx source file.
        chunk_size:    Maximum characters per prose chunk.
        chunk_overlap: Overlap in characters between prose chunks.

    Returns:
        List of Documents ready for embedding, carrying the same metadata
        contract as `md_splitter.split_md_file`.
    """
    document = open_document(docx_file)
    source_name = docx_file.name

    indexer = ChunkIndexer(source_name)
    sections = SectionStack()
    docs: list[Document] = []

    # Buffer of consecutive prose paragraphs, flushed on heading/table/size limit.
    buffer: list[str] = []
    buffer_start = 0
    buffer_end = 0
    buffer_len = 0

    def flush_prose() -> None:
        nonlocal buffer, buffer_start, buffer_end, buffer_len
        if buffer:
            docs.extend(prose_to_documents(
                "\n".join(buffer), indexer, sections.breadcrumb(),
                buffer_start, buffer_end, chunk_size, chunk_overlap,
            ))
        buffer = []
        buffer_start = buffer_end = buffer_len = 0

    for ordinal, block in enumerate(_iter_block_items(document), start=1):
        if isinstance(block, Paragraph):
            level = _heading_level(block)
            if level is not None:
                heading = clean_text(block.text)
                if heading:
                    flush_prose()
                    if _SERVICE_HEADING_RE.match(heading):
                        sections.reset(level)
                    else:
                        sections.push(level, heading)
                continue

            text = _paragraph_text(block)
            if not text:
                continue
            if not buffer:
                buffer_start = ordinal
            buffer.append(text)
            buffer_end = ordinal
            buffer_len += len(text) + 1
            if buffer_len >= chunk_size:
                flush_prose()
            continue

        # Table block
        flush_prose()
        breadcrumb = sections.breadcrumb()
        headers, data_rows = _parse_table(block)

        if headers and data_rows:
            columns = normalize_headers(headers)
            docs.append(Document(
                page_content=_render_table_text(columns, data_rows),
                metadata=indexer.meta(
                    breadcrumb, "table_full", ordinal, ordinal,
                    table_headers=json.dumps(columns, ensure_ascii=False),
                ),
            ))
            docs.extend(table_rows_to_documents(
                columns, data_rows, indexer, breadcrumb, ordinal, ordinal,
            ))
            logger.debug(
                f"[{source_name}] table '{breadcrumb[:60]}': "
                f"{len(data_rows)} rows, {len(headers)} columns"
            )
        else:
            raw = _raw_table_text(block)
            if raw.strip(" |\n"):
                docs.append(Document(
                    page_content=raw,
                    metadata=indexer.meta(breadcrumb, "table_raw", ordinal, ordinal),
                ))
                logger.debug(f"[{source_name}] table unparseable -> table_raw, bc='{breadcrumb[:60]}'")

    flush_prose()

    logger.debug(f"{source_name}: {len(docs)} chunks total")
    return docs
